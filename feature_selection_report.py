from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import os

def feature_selection_report(dataset):
    indexFile = 'data/datasets/{}/index.json'.format(dataset)
    reportFile = 'data/datasets/{}/feature_sel_report.docx'.format(dataset)

    with open(indexFile) as f:
        index = json.load(f)

    document = Document()
    #document.add_heading('{} Feature Selection'.format(dataset), 0)

    for _sym, files in index.items():
        corrPct = 'data/datasets/{}/correlation/{}_p1.png'.format(dataset, _sym)
        corrDif = 'data/datasets/{}/correlation/{}_d1.png'.format(dataset, _sym)
        featureImp = 'data/datasets/{}/feature_selection/{}.png'.format(dataset, _sym)
        featureImpPct = 'data/datasets/{}/feature_selection/{}_test30_pct.png'.format(dataset, _sym)
        featureImpDiff = 'data/datasets/{}/feature_selection/{}_test30_diff.png'.format(dataset, _sym)
        discussionFile = 'data/datasets/{}/feature_selection/{}_discussion.txt'.format(dataset, _sym)
        discussionText = 'discussion here'
        if os.path.exists(discussionFile):
            discussionText = open(discussionFile, 'r', encoding="utf8").read()

        document.add_heading(_sym, level=3)
        document.add_heading('Feature Importances', level=4)
        # document.add_picture(featureImp, width=Cm(16.8))
        # document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture(featureImpPct, width=Cm(16.8))
        document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_picture(featureImpDiff, width=Cm(16.8))
        # document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_heading('Correlation Matrix', level=4)
        document.add_picture(corrPct, width=Cm(10))
        document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_picture(corrDif, width=Cm(10))
        # document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_heading('Discussion', level=4)
        document.add_paragraph(discussionText)
        document.add_page_break()

    document.save(reportFile)

if __name__ == '__main__':
    feature_selection_report('ohlcv_coinmetrics')
